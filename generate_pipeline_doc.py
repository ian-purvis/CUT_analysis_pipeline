#!/usr/bin/env python3
"""
Generate CUT&RUN Pipeline documentation as a Word document (.docx).
Requires: pip install python-docx

Usage: python generate_pipeline_doc.py
Output: PIPELINE_DOCUMENTATION.docx (in cutrun_pipeline directory)
"""

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
    sys.exit(1)


def add_heading(doc, text, level=1):
    """Add a heading with consistent styling."""
    h = doc.add_heading(text, level=level)
    return h


def add_paragraph(doc, text, style=None):
    """Add a paragraph."""
    p = doc.add_paragraph(text, style=style)
    return p


def add_table_from_markdown(doc, lines):
    """Parse markdown table lines and add a Word table."""
    if not lines:
        return
    rows = []
    for line in lines:
        line = line.strip()
        if not line or line.startswith("|---"):
            continue
        cells = [c.strip() for c in re.split(r"\|", line) if c.strip()]
        if cells:
            rows.append(cells)
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            if j < len(table.rows[i].cells):
                table.rows[i].cells[j].text = cell
    doc.add_paragraph()


def parse_markdown_to_docx(md_path: Path, docx_path: Path):
    """Convert markdown file to Word document."""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.size = Pt(11)
    font.name = 'Calibri'

    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()

    lines = content.split("\n")
    i = 0
    in_table = False
    table_lines = []
    in_code = False
    code_lines = []

    while i < len(lines):
        line = lines[i]
        orig = line
        line_stripped = line.strip()

        # Code block
        if line_stripped.startswith("```"):
            if in_code:
                code_text = "\n".join(code_lines)
                p = doc.add_paragraph(code_text)
                p.style = "Normal"
                for run in p.runs:
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # Table
        if "|" in line_stripped and line_stripped.startswith("|"):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        else:
            if in_table and table_lines:
                add_table_from_markdown(doc, table_lines)
                table_lines = []
                in_table = False

        # Headings
        if line_stripped.startswith("# "):
            add_heading(doc, line_stripped[2:].strip(), level=0)
        elif line_stripped.startswith("## "):
            add_heading(doc, line_stripped[3:].strip(), level=1)
        elif line_stripped.startswith("### "):
            add_heading(doc, line_stripped[4:].strip(), level=2)
        elif line_stripped.startswith("#### "):
            add_heading(doc, line_stripped[5:].strip(), level=3)
        elif line_stripped.startswith("---"):
            pass  # Horizontal rule, skip
        elif line_stripped.startswith("*") or line_stripped.startswith("-"):
            # List item
            add_paragraph(doc, line_stripped)
        elif line_stripped:
            add_paragraph(doc, line_stripped)
        # Empty lines add spacing
        i += 1

    if in_table and table_lines:
        add_table_from_markdown(doc, table_lines)
    if in_code and code_lines:
        code_text = "\n".join(code_lines)
        p = doc.add_paragraph(code_text)
        for run in p.runs:
            run.font.name = "Consolas"
            run.font.size = Pt(9)

    doc.save(docx_path)
    print(f"Created: {docx_path}")


def main():
    script_dir = Path(__file__).resolve().parent
    md_path = script_dir / "PIPELINE_DOCUMENTATION.md"
    docx_path = script_dir / "PIPELINE_DOCUMENTATION.docx"

    if not md_path.exists():
        print(f"ERROR: {md_path} not found", file=sys.stderr)
        sys.exit(1)

    parse_markdown_to_docx(md_path, docx_path)


if __name__ == "__main__":
    main()
